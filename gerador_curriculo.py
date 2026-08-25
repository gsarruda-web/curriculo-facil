from pathlib import Path
from copy import deepcopy
from datetime import datetime
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE = BASE_DIR / "modelo_curriculo_base.docx"
OUT_DIR = BASE_DIR / "saidas"
OUT_DIR.mkdir(exist_ok=True)

MESES = {1:'Janeiro',2:'Fevereiro',3:'Março',4:'Abril',5:'Maio',6:'Junho',7:'Julho',8:'Agosto',9:'Setembro',10:'Outubro',11:'Novembro',12:'Dezembro'}

def _remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None

def _find_paragraph(doc, marker):
    for p in doc.paragraphs:
        if marker in p.text:
            return p
    return None

def _replace_text_everywhere(doc, replacements):
    for p in doc.paragraphs:
        full = ''.join(r.text for r in p.runs)
        changed = full
        for k,v in replacements.items():
            changed = changed.replace(k, v or '')
        if changed != full:
            # preserve formatting of first run, rebuild paragraph text
            if p.runs:
                p.runs[0].text = changed
                for r in p.runs[1:]: r.text = ''
            else:
                p.add_run(changed)

def _insert_lines_at_marker(doc, marker, lines, bullet_prefix=''):
    p = _find_paragraph(doc, marker)
    if not p: return
    if not lines:
        # remove heading immediately above marker plus marker itself
        prev_el = p._element.getprevious()
        if prev_el is not None:
            prev_el.getparent().remove(prev_el)
        _remove_paragraph(p)
        return
    for line in lines:
        np = p.insert_paragraph_before()
        np.paragraph_format.space_after = Pt(1)
        r = np.add_run((bullet_prefix if bullet_prefix else '') + line)
        r.font.name='Times New Roman'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); r.font.size=Pt(10.6)
    _remove_paragraph(p)

def _clean_text(s):
    return re.sub(r'\s+',' ',(s or '').strip())

def build_profile(data):
    skills = [_clean_text(x) for x in data.get('habilidades',[]) if _clean_text(x)]
    cargo = _clean_text(data.get('cargo_objetivo'))
    text = 'Profissional'
    if skills:
        text += ' com perfil marcado por ' + ', '.join(skills[:5]).lower()
    if cargo:
        text += f', com interesse em atuar na área de {cargo}'
    if data.get('experiencias'):
        text += ' e experiência prática compatível com o desenvolvimento das atividades profissionais'
    return text.rstrip('.') + '. Demonstra comprometimento, capacidade de aprendizagem e disposição para contribuir com a equipe e com os objetivos da organização.'

def build_objective(data):
    cargo = _clean_text(data.get('cargo_objetivo')) or 'a função pretendida'
    return (f'Atuar como {cargo}, desempenhando as atividades com responsabilidade, organização e comprometimento, '
            'contribuindo para a qualidade dos serviços da organização e para o desenvolvimento profissional contínuo.')

def format_education(item):
    tipo=_clean_text(item.get('tipo')); curso=_clean_text(item.get('curso')); inst=_clean_text(item.get('instituicao')); status=_clean_text(item.get('status')); ano=_clean_text(item.get('ano'))
    left = ' - '.join(x for x in [tipo, curso] if x)
    if inst: left += (' - ' if left else '') + inst
    tail = ' - '.join(x for x in [status, ano] if x)
    return left + ((' - ' + tail) if tail else '')

def format_course(item):
    nome=_clean_text(item.get('nome')); inst=_clean_text(item.get('instituicao')); ano=_clean_text(item.get('ano'))
    return nome + (f' - {inst}' if inst else '') + (f' - {ano}' if ano else '')

def format_experience(item):
    tipo=_clean_text(item.get('tipo')); local=_clean_text(item.get('empresa')); cargo=_clean_text(item.get('cargo')); periodo=_clean_text(item.get('periodo')); atividades=_clean_text(item.get('atividades'))
    parts=[]
    if local: parts.append(local)
    if cargo: parts.append(cargo)
    if periodo: parts.append(periodo)
    base=' - '.join(parts)
    if tipo and tipo.lower() not in ('emprego','experiência profissional'):
        base = f'{tipo}: ' + base
    if atividades:
        base += ('. ' if base else '') + atividades
    return base

def gerar_curriculo(data: dict, output_path: str | None = None):
    doc = Document(TEMPLATE)
    nome = _clean_text(data.get('nome_completo')).upper()
    cidade_uf = _clean_text(data.get('cidade_uf')).upper()
    cidade_data = _clean_text(data.get('cidade_data') or data.get('cidade_uf','').split('-')[0]).title()
    now = datetime.now()
    mes_ano = f'{MESES[now.month]} de {now.year}'
    endereco = _clean_text(data.get('endereco')).upper()
    if not endereco:
        p_end = _find_paragraph(doc, '[[ENDERECO]]')
        if p_end:
            _remove_paragraph(p_end)
    replacements = {
        '[[NOME_COMPLETO]]': nome,
        '[[ENDERECO]]': endereco,
        '[[CIDADE_UF]]': cidade_uf,
        '[[TELEFONE]]': _clean_text(data.get('telefone')),
        '[[EMAIL]]': _clean_text(data.get('email')),
        '[[CARGO_OBJETIVO]]': _clean_text(data.get('cargo_objetivo')),
        '[[PERFIL_PROFISSIONAL]]': _clean_text(data.get('perfil_profissional')) or build_profile(data),
        '[[OBJETIVO]]': _clean_text(data.get('objetivo')) or build_objective(data),
        '[[CIDADE_DATA]]': cidade_data,
        '[[MES_ANO]]': mes_ano,
    }
    _replace_text_everywhere(doc, replacements)
    _insert_lines_at_marker(doc, '[[CURSOS]]', [format_course(x) for x in data.get('cursos',[]) if format_course(x)])
    _insert_lines_at_marker(doc, '[[FORMACAO]]', [format_education(x) for x in data.get('formacao',[]) if format_education(x)])
    _insert_lines_at_marker(doc, '[[EXPERIENCIAS]]', [format_experience(x) for x in data.get('experiencias',[]) if format_experience(x)], bullet_prefix='- ')
    # clean possible blank contact lines with only labels
    for p in list(doc.paragraphs):
        if p.text.strip() in ('E-mail:', 'Área de atuação:', 'FONE:'):
            _remove_paragraph(p)
    slug = re.sub(r'[^A-Za-z0-9_-]+','_', nome.title()).strip('_') or 'curriculo'
    out = Path(output_path) if output_path else OUT_DIR / f'Curriculo_{slug}.docx'
    doc.save(out)
    return out
